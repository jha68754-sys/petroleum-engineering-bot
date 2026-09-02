"""
File processing module.

Handles extraction of text and tabular data from:
- PDF files (pypdf preferred, PyPDF2 fallback)
- DOCX files (python-docx)
- Excel files (openpyxl)
- CSV files (csv module)
- Image files (for vision AI analysis)

Implements automatic table detection, segmentation, and validation.
"""

from __future__ import annotations

import csv
import io
import logging
import os
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from constants import PVT_SECTION_HEADERS

logger = logging.getLogger("pvt_bot.services.file_processing")


def extract_pdf_text(file_path: str) -> str:
    """
    Extract text from a PDF file.

    Tries pypdf first (preferred), falls back to PyPDF2.
    Returns raw text with minimal formatting.

    Args:
        file_path: Path to the PDF file.

    Returns:
        Extracted text string, or empty string on failure.
    """
    text_parts: List[str] = []

    # Try pypdf first
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        result = "\n".join(text_parts)
        logger.info("PDF extracted via pypdf: %d chars from %d pages", len(result), len(reader.pages))
        return result
    except ImportError:
        logger.info("pypdf not available, trying PyPDF2")
    except Exception as exc:
        logger.warning("pypdf extraction failed: %s, trying PyPDF2", exc)

    # Fallback: PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        result = "\n".join(text_parts)
        logger.info("PDF extracted via PyPDF2: %d chars from %d pages", len(result), len(reader.pages))
        return result
    except ImportError:
        logger.error("Neither pypdf nor PyPDF2 is available")
        return ""
    except Exception as exc:
        logger.error("PyPDF2 extraction failed: %s", exc)
        return ""


def extract_docx_text(file_path: str) -> str:
    """
    Extract text from a DOCX file.

    Iterates over paragraphs and tables.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Extracted text string, or empty string on failure.
    """
    try:
        from docx import Document
        doc = Document(file_path)
        parts: List[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)

        result = "\n".join(parts)
        logger.info("DOCX extracted: %d chars", len(result))
        return result
    except ImportError:
        logger.error("python-docx not available")
        return ""
    except Exception as exc:
        logger.error("DOCX extraction failed: %s", exc)
        return ""


def extract_markdown_text(file_path: str) -> str:
    """Read UTF-8 Markdown used by the bot's Portable Case Snapshot."""
    try:
        with open(file_path, "r", encoding="utf-8-sig") as f:
            return f.read()
    except (OSError, UnicodeError) as exc:
        logger.error("Markdown extraction failed: %s", exc)
        return ""


def extract_csv_text(file_path: str) -> str:
    """
    Extract text from a CSV file with auto-detected delimiter.

    Detects comma, semicolon, or tab delimiters.
    Returns formatted text with headers.

    Args:
        file_path: Path to the CSV file.

    Returns:
        Formatted CSV content as text, or empty string on failure.
    """
    try:
        with open(file_path, "r", encoding="utf-8-sig") as f:
            raw = f.read()

        # Detect delimiter
        delimiter = ","
        if raw.count(";") > raw.count(","):
            delimiter = ";"
        elif raw.count("\t") > raw.count(","):
            delimiter = "\t"

        lines = raw.strip().split("\n")
        if not lines:
            return ""

        # Parse with csv module
        reader = csv.reader(io.StringIO(raw), delimiter=delimiter)
        rows = list(reader)

        if not rows:
            return ""

        # Format as text
        parts: List[str] = []
        for row in rows:
            parts.append(" | ".join(cell.strip() for cell in row))

        result = "\n".join(parts)
        logger.info("CSV extracted: %d rows, delimiter=%r", len(rows), delimiter)
        return f"__CSV__\n{result}"
    except Exception as exc:
        logger.error("CSV extraction failed: %s", exc)
        return ""


def extract_excel_text(file_path: str) -> str:
    """
    Extract text and tabular data from an Excel file (.xlsx/.xls).

    Reads all sheets, extracts data rows, and formats as text.
    Auto-detects PVT-related columns.

    Args:
        file_path: Path to the Excel file.

    Returns:
        Formatted Excel content as text, or empty string on failure.
    """
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        parts: List[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"\n=== Sheet: {sheet_name} ===")

            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                parts.append("(empty sheet)")
                continue

            for row in rows:
                cells = [
                    str(cell).strip() if cell is not None else ""
                    for cell in row
                ]
                parts.append(" | ".join(cells))

        wb.close()
        result = "\n".join(parts)
        logger.info("Excel extracted: %d chars from %d sheets", len(result), len(wb.sheetnames))
        return f"__EXCEL__\n{result}"
    except ImportError:
        logger.error("openpyxl not available for Excel extraction")
        return ""
    except Exception as exc:
        logger.error("Excel extraction failed: %s", exc)
        return ""


def extract_file_content(file_path: str, file_type: str) -> str:
    """
    Dispatch file extraction to the appropriate handler.

    Args:
        file_path: Path to the uploaded file.
        file_type: One of "pdf", "docx", "csv", "excel", "image".

    Returns:
        Extracted text content. For images, returns empty string
        (handled by vision AI separately).
    """
    if file_type == "image":
        return ""  # Images handled by vision AI

    extractors = {
        "pdf": extract_pdf_text,
        "docx": extract_docx_text,
        "markdown": extract_markdown_text,
        "csv": extract_csv_text,
        "excel": extract_excel_text,
    }

    extractor = extractors.get(file_type)
    if not extractor:
        logger.warning("Unknown file type: %s", file_type)
        return ""

    return extractor(file_path)


def segment_pdf_text(text: str, max_chars: int = 20000) -> List[str]:
    """
    Segment long PDF text into chunks at section boundaries.

    Tries to split at PVT section headers. Falls back to
    character-based splitting if no headers found.

    Args:
        text: The full extracted PDF text.
        max_chars: Maximum characters per chunk.

    Returns:
        List of text chunks.
    """
    if len(text) <= max_chars:
        return [text]

    # Try splitting at PVT section headers
    segments: List[str] = []
    current = ""

    for line in text.split("\n"):
        if len(current) + len(line) > max_chars and current:
            segments.append(current)
            current = ""

        # Check if this line starts a new section
        is_section = any(
            header.lower() in line.lower()
            for header in PVT_SECTION_HEADERS
        )

        if is_section and current:
            segments.append(current)
            current = ""

        current += line + "\n"

    if current:
        segments.append(current)

    logger.info("PDF segmented into %d chunks", len(segments))
    return segments


def format_segmented_context(segments: List[str]) -> str:
    """
    Format segmented text for AI context injection.

    Args:
        segments: List of text segments.

    Returns:
        Formatted context string with segment markers.
    """
    parts: List[str] = []
    for i, seg in enumerate(segments):
        parts.append(f"--- Segment {i + 1} ---\n{seg.strip()}")
    return "\n".join(parts)


def detect_file_type(filename: str) -> str:
    """
    Detect file type from extension.

    Args:
        filename: The filename (with extension).

    Returns:
        One of "pdf", "docx", "csv", "excel", "image", or "unknown".
    """
    ext = Path(filename).suffix.lower()
    type_map = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".md": "markdown",
        ".markdown": "markdown",
        ".csv": "csv",
        ".xlsx": "excel",
        ".xls": "excel",
        ".png": "image",
        ".jpg": "image",
        ".jpeg": "image",
        ".webp": "image",
    }
    return type_map.get(ext, "unknown")


def save_uploaded_file(file_bytes: bytes, filename: str) -> Optional[str]:
    """
    Save uploaded file bytes to a temporary file.

    Args:
        file_bytes: Raw file content.
        filename: Original filename.

    Returns:
        Path to the saved temp file, or None on failure.
    """
    try:
        ext = Path(filename).suffix.lower()
        fd, tmp_path = tempfile.mkstemp(suffix=ext, prefix="pvt_upload_")
        with os.fdopen(fd, "wb") as f:
            f.write(file_bytes)
        logger.info("File saved: %s (%d bytes)", tmp_path, len(file_bytes))
        return tmp_path
    except Exception as exc:
        logger.error("Failed to save uploaded file: %s", exc)
        return None
